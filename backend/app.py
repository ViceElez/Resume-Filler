import ollama
from docx import Document
from flask import Flask, request, jsonify
from flask_cors import CORS
import os
from groq import Groq
from dotenv import load_dotenv
from pathlib import Path

app = Flask(__name__)
CORS(app)
env_path = Path(__file__).resolve().parent.parent / '.env'
load_dotenv(dotenv_path=env_path)
api_key = os.getenv("API_KEY")
model="llama-3.3-70b-versatile"

client = Groq(
    api_key=api_key
)

@app.route('/upload', methods=['POST'])
def upload_files():
    user_resume = request.files.get('user-resume-upload-word')
    company_resume = request.files.get('company-resume-upload-word')
    user_resume_docx = Document(user_resume)
    company_resume_docx = Document(company_resume)

    user_resume_text = "\n".join([p.text for p in user_resume_docx.paragraphs])
    company_resume_text = "\n".join([p.text for p in company_resume_docx.paragraphs])

    prompt=f""" Rules:
        - Only fill fields if the field name exactly matches a key in userResume.
        - If a field does not exist in userResume, write exactly: "No relative data for this".
        - Keep the formatting exactly as in companyResume.
        - Do NOT add extra fields, explanations, or text.
        - Do NOT change the order of fields in companyResume.
        - Do NOT type filled from userResume.

        Input:

        userResume:
        {user_resume_text}

        companyResume:
        {company_resume_text}

        Output:
        """

    response = client.chat.completions.create(
    messages=[
        {
            "role": "user",
            "content": prompt,
        }
    ],
    model=model,
)

    return jsonify({"result": response.choices[0].message.content})

if __name__ == '__main__':
    app.run(debug=True)